"""
Modulo di trasformazione: applica le sostituzioni ai file originali
in base alle decisioni di review dell'utente.
"""

import logging
import shutil
from pathlib import Path
from typing import Dict, List, Optional

from app.models.schemas import Finding, ReviewAction
from app.parsers.base import ParseResult

logger = logging.getLogger(__name__)


def _build_substitution_map(findings: List[Finding]) -> Dict[str, str]:
    """
    Costruisce una mappa {valore_originale: pseudonimo_finale}
    in base alle decisioni di review.
    I finding con azione REJECT vengono esclusi.
    """
    sub_map: Dict[str, str] = {}
    for finding in findings:
        if finding.review_action == ReviewAction.REJECT:
            continue
        final = finding.final_pseudonym
        if finding.original_value not in sub_map:
            sub_map[finding.original_value] = final
    return sub_map


def _apply_substitutions_to_text(text: str, sub_map: Dict[str, str]) -> str:
    """
    Applica le sostituzioni a una stringa di testo.
    Ordina le chiavi per lunghezza decrescente per evitare sostituzioni parziali.
    """
    if not sub_map:
        return text

    # Ordina per lunghezza decrescente (sostituisce prima le stringhe più lunghe)
    sorted_keys = sorted(sub_map.keys(), key=len, reverse=True)
    result = text
    for original in sorted_keys:
        pseudonym = sub_map[original]
        # Usa re.sub per una sostituzione case-sensitive precisa
        result = result.replace(original, pseudonym)
    return result


def transform_text_file(
    original_path: Path,
    output_path: Path,
    findings: List[Finding],
) -> List[str]:
    """Trasforma un file di testo semplice (.txt, .md, .csv)."""
    warnings = []
    sub_map = _build_substitution_map(findings)

    try:
        try:
            content = original_path.read_text(encoding="utf-8")
        except UnicodeDecodeError:
            content = original_path.read_text(encoding="latin-1")
            warnings.append("File letto con codifica latin-1.")

        transformed = _apply_substitutions_to_text(content, sub_map)
        output_path.write_text(transformed, encoding="utf-8")
    except Exception as e:
        warnings.append(f"Errore durante la trasformazione del file di testo: {e}")

    return warnings


def transform_docx_file(
    original_path: Path,
    output_path: Path,
    findings: List[Finding],
) -> List[str]:
    """Trasforma un file .docx."""
    warnings = []
    sub_map = _build_substitution_map(findings)

    try:
        from docx import Document

        doc = Document(str(original_path))

        # Trasforma paragrafi
        for para in doc.paragraphs:
            if para.text.strip():
                for run in para.runs:
                    run.text = _apply_substitutions_to_text(run.text, sub_map)

        # Trasforma tabelle
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        for run in para.runs:
                            run.text = _apply_substitutions_to_text(run.text, sub_map)

        # Trasforma header e footer
        for section in doc.sections:
            for part in [section.header, section.footer]:
                if part:
                    for para in part.paragraphs:
                        for run in para.runs:
                            run.text = _apply_substitutions_to_text(run.text, sub_map)

        doc.save(str(output_path))
        warnings.append("LIMITE MVP: Commenti, note a piè di pagina e caselle di testo non sono stati processati.")

    except Exception as e:
        warnings.append(f"Errore durante la trasformazione del file DOCX: {e}")

    return warnings


def transform_xlsx_file(
    original_path: Path,
    output_path: Path,
    findings: List[Finding],
) -> List[str]:
    """Trasforma un file .xlsx (solo celle testuali, formule intatte)."""
    warnings = []
    sub_map = _build_substitution_map(findings)

    try:
        import openpyxl

        wb = openpyxl.load_workbook(str(original_path))

        modified_cells = 0
        for sheet_name in wb.sheetnames:
            ws = wb[sheet_name]
            for row in ws.iter_rows():
                for cell in row:
                    # Non toccare le formule
                    if isinstance(cell.value, str) and cell.value.startswith("="):
                        continue
                    if isinstance(cell.value, str) and cell.value.strip():
                        new_value = _apply_substitutions_to_text(cell.value, sub_map)
                        if new_value != cell.value:
                            cell.value = new_value
                            modified_cells += 1

        wb.save(str(output_path))
        warnings.append(f"Modificate {modified_cells} celle testuali. Le formule sono state preservate.")

    except Exception as e:
        warnings.append(f"Errore durante la trasformazione del file XLSX: {e}")

    return warnings


# ─── PDF Transformation Helpers ───────────────────────────────────────────────


def _extract_pdf_text_by_page(original_path: Path) -> tuple:
    """
    Estrae testo da PDF, pagina per pagina.
    Gestisce PDF cifrati, non testuali.
    Restituisce (pages_text, warnings, is_encrypted, is_empty).
    """
    from pypdf import PdfReader

    warnings = []
    pages_text = []

    reader = PdfReader(str(original_path))

    if reader.is_encrypted:
        return [], ["PDF è cifrato/protetto."], True, False

    for page_num, page in enumerate(reader.pages, start=1):
        try:
            text = page.extract_text() or ""
            pages_text.append(text)
        except Exception as pe:
            pages_text.append("")
            warnings.append(f"Errore estrazione pagina {page_num}: {pe}")

    is_empty = not any(t.strip() for t in pages_text)
    return pages_text, warnings, False, is_empty


def _rebuild_pdf_from_pages(pages_pseudo: List[str], output_path: Path) -> List[str]:
    """
    Ricostruisce PDF da pagine di testo pseudonimizzato.
    Tenta fpdf2, fallback a reportlab.
    Restituisce lista di warning.
    """
    warnings = []

    # Tentativo primario: fpdf2
    try:
        from fpdf import FPDF

        pdf = FPDF()
        pdf.set_auto_page_break(auto=True, margin=15)
        pdf.set_margins(15, 15, 15)
        pdf.set_title("")
        pdf.set_author("")
        pdf.set_creator("Local Pseudonymization Tool")
        pdf.set_subject("")

        for page_text in pages_pseudo:
            pdf.add_page()
            pdf.set_font("Helvetica", size=10)
            for line in page_text.splitlines():
                safe_line = line.encode("latin-1", errors="replace").decode("latin-1")
                pdf.multi_cell(0, 5, safe_line)

        pdf.output(str(output_path))
        warnings.append(
            "PDF rebuild completato (fpdf2). Il layout potrebbe essere differente dall'originale."
        )
        return warnings

    except Exception as fpdf_err:
        warnings.append(f"fpdf2 fallito ({fpdf_err}), tentativo reportlab...")

    # Fallback: reportlab
    try:
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.styles import getSampleStyleSheet
        from reportlab.lib.units import mm
        from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer

        styles = getSampleStyleSheet()
        story = []
        for page_text in pages_pseudo:
            for line in page_text.splitlines():
                safe = line.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
                if safe.strip():
                    story.append(Paragraph(safe, styles["Normal"]))
                else:
                    story.append(Spacer(1, 3 * mm))
            story.append(Spacer(1, 10 * mm))

        doc_rl = SimpleDocTemplate(str(output_path), pagesize=A4)
        doc_rl.build(story)
        warnings.append(
            "PDF rebuild completato (reportlab). Il layout potrebbe essere differente dall'originale."
        )
        return warnings

    except Exception as rl_err:
        warnings.append(f"Errore reportlab: {rl_err}")
        return warnings


# ─── transform_pdf_file ──────────────────────────────────────────────────────


def transform_pdf_file(
    original_path: Path,
    output_path: Path,
    findings: List[Finding],
    strict: bool = False,
) -> List[str]:
    """
    Trasforma PDF testuale in PDF pseudonimizzato (PDF→PDF).
    Estrategia: testo → sostituzioni → rebuild.
    Layout non preservato. Se cifrato/non-testuale: copia + warning.
    """
    warnings = []
    sub_map = _build_substitution_map(findings)

    # Assicura output .pdf
    if output_path.suffix.lower() != ".pdf":
        output_path = output_path.with_suffix(".pdf")

    try:
        # Step 1: Estrai testo per pagina
        pages_text, extract_warnings, is_encrypted, is_empty = _extract_pdf_text_by_page(original_path)
        warnings.extend(extract_warnings)

        # Step 2: Gestisci PDF cifrato o non-testuale
        if is_encrypted:
            warnings.append(
                "PDF cifrato/protetto: impossibile estrarre testo. " "File NON pseudonimizzato. Etichetta: NOT_SAFE."
            )
            shutil.copy2(str(original_path), str(output_path))
            return warnings

        if is_empty:
            warnings.append(
                "PDF non testuale (scansionato?). File NON pseudonimizzato. " "Etichetta: SAFE_WITH_WARNINGS."
            )
            shutil.copy2(str(original_path), str(output_path))
            return warnings

        # Step 3: Applica sostituzioni
        pages_pseudo = [_apply_substitutions_to_text(t, sub_map) for t in pages_text]

        # Step 4: Ricostruisci PDF
        rebuild_warnings = _rebuild_pdf_from_pages(pages_pseudo, output_path)
        warnings.extend(rebuild_warnings)

        if not rebuild_warnings or "fallito" in rebuild_warnings[0]:
            shutil.copy2(str(original_path), str(output_path))

    except Exception as e:
        warnings.append(f"Errore trasformazione PDF: {e}")
        shutil.copy2(str(original_path), str(output_path))

    return warnings


def transform_image_file(
    original_path: Path,
    output_path: Path,
    findings: List[Finding],
    parse_result: Optional[ParseResult] = None,
) -> List[str]:
    """
    Applica la redazione visuale (box di oscuramento) sull'immagine
    per le entità trovate tramite OCR.
    """
    warnings = []

    try:
        from PIL import Image, ImageDraw

        # Usa l'immagine già pulita dall'EXIF se disponibile
        source_path = parse_result.image_path if parse_result and parse_result.image_path else original_path
        img = Image.open(str(source_path))
        draw = ImageDraw.Draw(img)

        redacted_count = 0
        for finding in findings:
            if finding.review_action == ReviewAction.REJECT:
                continue
            bbox = finding.location.bbox
            if bbox and len(bbox) == 4:
                x, y, w, h = bbox
                # Disegna un rettangolo nero di redazione
                draw.rectangle([x, y, x + w, y + h], fill="black", outline="black")
                redacted_count += 1

        img.save(str(output_path))
        warnings.append(f"Redazione visuale applicata: {redacted_count} aree oscurate.")

        if redacted_count == 0 and findings:
            warnings.append(
                "ATTENZIONE: Nessuna area di redazione visuale applicata. "
                "Le entità trovate potrebbero non avere coordinate bounding box precise."
            )

    except Exception as e:
        warnings.append(f"Errore durante la redazione visuale dell'immagine: {e}")

    return warnings


def transform_file(
    original_path: Path,
    output_dir: Path,
    findings: List[Finding],
    parse_result: Optional[ParseResult] = None,
) -> tuple[Path, List[str]]:
    """
    Dispatcher principale: seleziona la funzione di trasformazione appropriata
    in base all'estensione del file.

    Restituisce (output_path, warnings).
    """
    ext = original_path.suffix.lower()
    output_path = output_dir / original_path.name
    warnings: List[str] = []

    if ext in (".txt", ".md", ".csv"):
        warnings = transform_text_file(original_path, output_path, findings)

    elif ext == ".docx":
        warnings = transform_docx_file(original_path, output_path, findings)

    elif ext == ".xlsx":
        warnings = transform_xlsx_file(original_path, output_path, findings)

    elif ext == ".pdf":
        # Assicura sempre estensione .pdf nell'output
        output_path = output_dir / (original_path.stem + ".pdf")
        warnings = transform_pdf_file(original_path, output_path, findings)

    elif ext in (".jpg", ".jpeg", ".png"):
        warnings = transform_image_file(original_path, output_path, findings, parse_result)

    else:
        # Formato non supportato: copia il file originale con warning
        shutil.copy2(str(original_path), str(output_path))
        warnings.append(f"Formato '{ext}' non supportato per la trasformazione. File copiato senza modifiche.")

    return output_path, warnings


def apply_pseudonyms_to_text(
    text: str,
    findings: List[Finding],
) -> tuple:
    """
    Applica le sostituzioni a una stringa di testo puro (per console/clipboard).
    Restituisce (testo_pseudonimizzato, numero_sostituzioni_applicate).
    """
    sub_map = _build_substitution_map(findings)
    if not sub_map:
        return text, 0

    result = text
    applied = 0
    sorted_keys = sorted(sub_map.keys(), key=len, reverse=True)
    for original in sorted_keys:
        pseudonym = sub_map[original]
        if original in result:
            result = result.replace(original, pseudonym)
            applied += 1

    return result, applied
