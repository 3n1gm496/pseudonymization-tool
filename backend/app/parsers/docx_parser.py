"""
Parser per file .docx (Microsoft Word).
Estrae testo da paragrafi, tabelle, header e footer.
"""
from pathlib import Path
from typing import List

from app.parsers.base import BaseParser, ParseResult, TextChunk


class DocxParser(BaseParser):
    """Parser per file Microsoft Word (.docx)."""

    @property
    def supported_extensions(self) -> List[str]:
        return [".docx"]

    def supports_streaming(self) -> bool:
        return True

    def parse_stream(self, file_path: Path, chunk_size: int = 1000):
        """Parse DOCX in streaming, processando elemento per elemento."""
        from docx import Document
        try:
            doc = Document(str(file_path))
            for i, para in enumerate(doc.paragraphs):
                if para.text.strip():
                    yield TextChunk(
                        text=para.text,
                        source_ref=f"paragrafo {i + 1}",
                        line_number=i + 1,
                    )

            for t_idx, table in enumerate(doc.tables):
                for r_idx, row in enumerate(table.rows):
                    for c_idx, cell in enumerate(row.cells):
                        if cell.text.strip():
                            yield TextChunk(
                                text=cell.text,
                                source_ref=f"tabella {t_idx + 1}, riga {r_idx + 1}, col {c_idx + 1}",
                            )

            for s_idx, section in enumerate(doc.sections):
                for part_name, part in [("header", section.header), ("footer", section.footer)]:
                    if part is not None:
                        for para in part.paragraphs:
                            if para.text.strip():
                                yield TextChunk(
                                    text=para.text,
                                    source_ref=f"sezione {s_idx + 1} {part_name}",
                                )
        except Exception:
            return

    def parse(self, file_path: Path) -> ParseResult:
        result = ParseResult(file_path=file_path)
        try:
            from docx import Document
            doc = Document(str(file_path))

            # Paragrafi del corpo principale
            for i, para in enumerate(doc.paragraphs):
                if para.text.strip():
                    result.chunks.append(
                        TextChunk(
                            text=para.text,
                            source_ref=f"paragrafo {i + 1}",
                            line_number=i + 1,
                        )
                    )

            # Testo nelle tabelle
            for t_idx, table in enumerate(doc.tables):
                for r_idx, row in enumerate(table.rows):
                    for c_idx, cell in enumerate(row.cells):
                        if cell.text.strip():
                            result.chunks.append(
                                TextChunk(
                                    text=cell.text,
                                    source_ref=f"tabella {t_idx + 1}, riga {r_idx + 1}, col {c_idx + 1}",
                                )
                            )

            # Header e footer di ogni sezione
            for s_idx, section in enumerate(doc.sections):
                for part_name, part in [("header", section.header), ("footer", section.footer)]:
                    if part is not None:
                        for para in part.paragraphs:
                            if para.text.strip():
                                result.chunks.append(
                                    TextChunk(
                                        text=para.text,
                                        source_ref=f"sezione {s_idx + 1} {part_name}",
                                    )
                                )

            if not result.chunks:
                result.warnings.append("Il documento non contiene testo estraibile nei paragrafi principali.")

            result.warnings.append(
                "LIMITE MVP: Commenti, note a piè di pagina, caselle di testo e macro non vengono processati."
            )

        except Exception as e:
            result.success = False
            result.error_message = f"Errore durante il parsing del file DOCX: {e}"

        return result
