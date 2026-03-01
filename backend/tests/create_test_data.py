"""
Script per creare dataset di test sintetici realistici.
Genera file di test per tutti i formati supportati.
"""

import os
import sys
from pathlib import Path

# Aggiungi il backend al path
sys.path.insert(0, str(Path(__file__).parent.parent))

TEST_DATA_DIR = Path(__file__).parent / "test_data"
TEST_DATA_DIR.mkdir(exist_ok=True)


def create_txt_test():
    """File TXT con vari tipi di entità sensibili."""
    content = """# Log di Sistema — Incidente di Sicurezza
# Data: 2024-03-15 09:23:41

[ALERT] Tentativo di accesso non autorizzato rilevato.
Utente: mario.rossi@ente.gov.it
IP Sorgente: 10.24.8.15
IP Destinazione: 192.168.1.100
Hostname: srv-dc-01.ente.local
URL Accesso: https://intranet.ente.gov.it/admin/login

[INFO] Analisi log firewall completata.
Connessioni sospette da: 203.0.113.42
Verso: vpn.ente.local (10.0.0.1)
Utente coinvolto: luigi.ferrari@ente.gov.it
Codice Fiscale: RSSMRA80A01H501A
Telefono: +39 333 1234567

[WARNING] Accesso multiplo fallito per l'account:
Username: @m.rossi
Email alternativa: m.rossi@comune.esempio.it
Partita IVA fornitore: 12345678901

[INFO] IPv6 rilevato nel traffico:
Indirizzo: 2001:0db8:85a3:0000:0000:8a2e:0370:7334
Hostname: server-prod-01.ente.local

[RESOLVED] Incidente chiuso. Contatto: anna.bianchi@ente.gov.it
Telefono ufficio: 06 1234 5678
"""
    (TEST_DATA_DIR / "test_log.txt").write_text(content, encoding="utf-8")
    print("✓ Creato test_log.txt")


def create_md_test():
    """File Markdown con dati sensibili."""
    content = """# Policy di Sicurezza — Bozza Riservata

## 1. Contatti Responsabili

Il responsabile della sicurezza è **Giovanni Conti** (giovanni.conti@ente.gov.it).
Per emergenze contattare: +39 06 9876 5432

## 2. Infrastruttura

I server principali sono:
- `srv-dc-01.ente.local` (10.0.0.1) — Domain Controller
- `srv-mail-01.ente.local` (10.0.0.2) — Mail Server
- `srv-vpn-01.ente.local` (10.0.0.3) — VPN Gateway

## 3. Accessi

Gli accessi sono gestiti tramite il portale: https://intranet.ente.gov.it/accessi

Utenti autorizzati:
- mario.rossi@ente.gov.it (CF: RSSMRA80A01H501A)
- anna.bianchi@ente.gov.it

## 4. Note

Progetto di riferimento: PROJ-2024-001
Unità organizzativa: UO-SICUREZZA
"""
    (TEST_DATA_DIR / "test_policy.md").write_text(content, encoding="utf-8")
    print("✓ Creato test_policy.md")


def create_csv_test():
    """File CSV con dati anagrafici."""
    content = """nome,cognome,email,telefono,codice_fiscale,ip_workstation
Mario,Rossi,mario.rossi@ente.gov.it,+39 333 1234567,RSSMRA80A01H501A,10.24.1.15
Luigi,Ferrari,luigi.ferrari@ente.gov.it,06 1234 5678,FRRLGU75B12F205X,10.24.1.16
Anna,Bianchi,anna.bianchi@ente.gov.it,+39 347 9876543,BNCNNA82C41H501Z,10.24.1.17
Giovanni,Conti,giovanni.conti@ente.gov.it,06 9876 5432,CNTGNN70D15H501W,10.24.1.18
"""
    (TEST_DATA_DIR / "test_users.csv").write_text(content, encoding="utf-8")
    print("✓ Creato test_users.csv")


def create_docx_test():
    """File DOCX con testo sensibile."""
    try:
        from docx import Document

        doc = Document()

        doc.add_heading("Relazione Tecnica — Riservata", 0)
        doc.add_paragraph(
            "Autore: Ing. Mario Rossi (mario.rossi@ente.gov.it)\n" "Data: 15 Marzo 2024\n" "Progetto: PROJ-2024-001"
        )

        doc.add_heading("1. Sommario Esecutivo", 1)
        doc.add_paragraph(
            "Durante l'analisi del sistema srv-dc-01.ente.local (IP: 10.24.8.1) "
            "è stato rilevato un accesso anomalo dall'indirizzo 203.0.113.42. "
            "Il responsabile Luigi Ferrari (luigi.ferrari@ente.gov.it, CF: FRRLGU75B12F205X) "
            "è stato notificato al numero +39 333 9876543."
        )

        doc.add_heading("2. Dettagli Tecnici", 1)
        table = doc.add_table(rows=1, cols=3)
        table.style = "Table Grid"
        hdr = table.rows[0].cells
        hdr[0].text = "Sistema"
        hdr[1].text = "IP"
        hdr[2].text = "Responsabile"

        data = [
            ("srv-dc-01.ente.local", "10.0.0.1", "mario.rossi@ente.gov.it"),
            ("srv-mail-01.ente.local", "10.0.0.2", "anna.bianchi@ente.gov.it"),
        ]
        for sys_name, ip, resp in data:
            row = table.add_row().cells
            row[0].text = sys_name
            row[1].text = ip
            row[2].text = resp

        doc.save(str(TEST_DATA_DIR / "test_report.docx"))
        print("✓ Creato test_report.docx")
    except Exception as e:
        print(f"✗ Errore creazione DOCX: {e}")


def create_xlsx_test():
    """File XLSX con testo e formule."""
    try:
        import openpyxl

        wb = openpyxl.Workbook()
        ws = wb.active
        ws.title = "Dati Utenti"

        # Header
        ws.append(["Nome", "Email", "IP", "Codice Fiscale", "Stipendio", "Bonus", "Totale"])

        # Dati
        ws.append(["Mario Rossi", "mario.rossi@ente.gov.it", "10.24.1.15", "RSSMRA80A01H501A", 2500, 500, "=E2+F2"])
        ws.append(["Luigi Ferrari", "luigi.ferrari@ente.gov.it", "10.24.1.16", "FRRLGU75B12F205X", 2800, 300, "=E3+F3"])
        ws.append(["Anna Bianchi", "anna.bianchi@ente.gov.it", "10.24.1.17", "BNCNNA82C41H501Z", 2600, 400, "=E4+F4"])

        # Secondo foglio con note
        ws2 = wb.create_sheet("Note")
        ws2["A1"] = "Contatto amministrativo: giovanni.conti@ente.gov.it"
        ws2["A2"] = "Server HR: srv-hr-01.ente.local (10.0.0.5)"
        ws2["A3"] = "Totale stipendi"
        ws2["B3"] = "=SUM(Dati Utenti!E2:E4)"

        wb.save(str(TEST_DATA_DIR / "test_data.xlsx"))
        print("✓ Creato test_data.xlsx")
    except Exception as e:
        print(f"✗ Errore creazione XLSX: {e}")


def create_pdf_test():
    """File PDF testuale."""
    try:
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.styles import getSampleStyleSheet
        from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer

        doc = SimpleDocTemplate(str(TEST_DATA_DIR / "test_document.pdf"), pagesize=A4)
        styles = getSampleStyleSheet()
        story = []

        story.append(Paragraph("Documento Riservato — Analisi di Sicurezza", styles["Title"]))
        story.append(Spacer(1, 12))
        story.append(
            Paragraph(
                "Responsabile: Dott. Mario Rossi (mario.rossi@ente.gov.it) — CF: RSSMRA80A01H501A", styles["Normal"]
            )
        )
        story.append(Spacer(1, 12))
        story.append(
            Paragraph(
                "Durante l'analisi del server srv-dc-01.ente.local (10.24.8.1) è stato rilevato "
                "traffico anomalo verso 203.0.113.42. Il firewall ha bloccato 15 tentativi di "
                "connessione verso https://malicious.example.com/payload.",
                styles["Normal"],
            )
        )
        story.append(Spacer(1, 12))
        story.append(Paragraph("Contatti: luigi.ferrari@ente.gov.it, tel. +39 06 1234 5678", styles["Normal"]))

        doc.build(story)
        print("✓ Creato test_document.pdf")
    except ImportError:
        # Fallback: crea un PDF minimale senza reportlab
        pdf_content = b"""%PDF-1.4
1 0 obj
<< /Type /Catalog /Pages 2 0 R >>
endobj
2 0 obj
<< /Type /Pages /Kids [3 0 R] /Count 1 >>
endobj
3 0 obj
<< /Type /Page /Parent 2 0 R /MediaBox [0 0 612 792]
/Contents 4 0 R /Resources << /Font << /F1 5 0 R >> >> >>
endobj
4 0 obj
<< /Length 200 >>
stream
BT
/F1 12 Tf
50 750 Td
(Documento Riservato - mario.rossi@ente.gov.it) Tj
0 -20 Td
(IP: 10.24.8.15 - CF: RSSMRA80A01H501A) Tj
ET
endstream
endobj
5 0 obj
<< /Type /Font /Subtype /Type1 /BaseFont /Helvetica >>
endobj
xref
0 6
0000000000 65535 f
0000000009 00000 n
0000000058 00000 n
0000000115 00000 n
0000000274 00000 n
0000000528 00000 n
trailer
<< /Size 6 /Root 1 0 R >>
startxref
610
%%EOF"""
        (TEST_DATA_DIR / "test_document.pdf").write_bytes(pdf_content)
        print("✓ Creato test_document.pdf (minimale)")
    except Exception as e:
        print(f"✗ Errore creazione PDF: {e}")


def create_image_test():
    """Immagine PNG con testo sensibile per test OCR."""
    try:
        import struct

        from PIL import Image, ImageDraw, ImageFont

        # Crea un'immagine bianca con testo
        img = Image.new("RGB", (800, 400), color="white")
        draw = ImageDraw.Draw(img)

        # Testo con dati sensibili
        lines = [
            "SISTEMA: srv-dc-01.ente.local",
            "IP: 10.24.8.15",
            "UTENTE: mario.rossi@ente.gov.it",
            "CF: RSSMRA80A01H501A",
            "STATO: ACCESSO FALLITO",
            "URL: https://intranet.ente.gov.it/admin",
        ]

        y = 40
        for line in lines:
            draw.text((40, y), line, fill="black")
            y += 50

        # Aggiungi metadati EXIF finti (per testare lo stripping)
        img_path = TEST_DATA_DIR / "test_screenshot.png"
        img.save(str(img_path))

        # Aggiungi metadati EXIF usando piexif se disponibile
        try:
            import piexif

            exif_dict = {
                "0th": {
                    piexif.ImageIFD.Make: b"TestCamera",
                    piexif.ImageIFD.Model: b"TestModel",
                    piexif.ImageIFD.Artist: b"mario.rossi@ente.gov.it",
                },
            }
            exif_bytes = piexif.dump(exif_dict)
            img_jpeg_path = TEST_DATA_DIR / "test_screenshot_exif.jpg"
            img.save(str(img_jpeg_path), "JPEG", exif=exif_bytes)
            print("✓ Creato test_screenshot_exif.jpg (con EXIF)")
        except ImportError:
            pass

        print("✓ Creato test_screenshot.png")
    except Exception as e:
        print(f"✗ Errore creazione immagine: {e}")


def create_pdf_non_textual():
    """PDF non testuale (basato su immagine) per testare il warning."""
    try:
        import io

        from PIL import Image

        # Crea un'immagine e salvala come PDF
        img = Image.new("RGB", (400, 200), color="white")
        draw = img if True else None

        from PIL import ImageDraw

        draw = ImageDraw.Draw(img)
        draw.text((20, 80), "Documento scansionato - non testuale", fill="black")

        img_path = TEST_DATA_DIR / "test_scanned.pdf"
        img.save(str(img_path), "PDF")
        print("✓ Creato test_scanned.pdf (non testuale, per test warning)")
    except Exception as e:
        print(f"✗ Errore creazione PDF non testuale: {e}")


if __name__ == "__main__":
    print("Creazione dataset di test sintetici...")
    create_txt_test()
    create_md_test()
    create_csv_test()
    create_docx_test()
    create_xlsx_test()
    create_pdf_test()
    create_image_test()
    create_pdf_non_textual()
    print(f"\nDataset creato in: {TEST_DATA_DIR}")
    print("File creati:")
    for f in sorted(TEST_DATA_DIR.iterdir()):
        print(f"  {f.name} ({f.stat().st_size} bytes)")
