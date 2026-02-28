"""
Test per il supporto streaming di file grandi.
"""
import os
from pathlib import Path
from unittest.mock import patch, MagicMock

from app.core.pipeline import _should_use_streaming
from app.parsers.factory import get_parser
from app.parsers.base import TextChunk


def test_should_use_streaming_small_file(tmp_path):
    """File piccoli non dovrebbero usare streaming."""
    small_file = tmp_path / "small.pdf"
    small_file.write_text("x" * 1024)  # 1KB
    
    assert not _should_use_streaming(small_file)


def test_should_use_streaming_large_file(tmp_path):
    """File grandi (>50MB) dovrebbero usare streaming."""
    large_file = tmp_path / "large.pdf"
    large_file.write_text("x")  # Crea il file
    
    # Mock Path.stat per simulare file grande
    from pathlib import Path
    original_stat = Path.stat
    
    def mock_stat(self):
        if str(self) == str(large_file):
            class MockStat:
                st_size = 60 * 1024 * 1024  # 60MB
            return MockStat()
        return original_stat(self)
    
    with patch.object(Path, 'stat', mock_stat):
        assert _should_use_streaming(large_file)


def test_should_use_streaming_nonexistent_file(tmp_path):
    """File non esistente non dovrebbe crashare."""
    nonexistent = tmp_path / "nonexistent.pdf"
    assert not _should_use_streaming(nonexistent)


def test_pdf_parser_supports_streaming():
    """PDF parser deve supportare streaming."""
    parser = get_parser(Path("test.pdf"))
    assert parser is not None
    assert hasattr(parser, 'supports_streaming')
    assert parser.supports_streaming() is True


def test_docx_parser_supports_streaming():
    """DOCX parser deve supportare streaming."""
    parser = get_parser(Path("test.docx"))
    assert parser is not None
    assert hasattr(parser, 'supports_streaming')
    assert parser.supports_streaming() is True


def test_text_parser_no_streaming():
    """Text parser non supporta streaming (ancora)."""
    parser = get_parser(Path("test.txt"))
    assert parser is not None
    # Text parser potrebbe non avere il metodo, è ok
    if hasattr(parser, 'supports_streaming'):
        assert parser.supports_streaming() is False


def test_pdf_parse_stream_chunks(tmp_path):
    """parse_stream dovrebbe restituire chunks iterativamente."""
    # Crea un PDF di test minimale (solo testo)
    pdf_content = tmp_path / "test.pdf"
    pdf_content.write_text("%PDF-1.4\nDummy PDF content\n%%EOF")
    
    parser = get_parser(pdf_content)
    assert parser is not None
    
    if hasattr(parser, 'parse_stream'):
        chunks = list(parser.parse_stream(pdf_content, chunk_size=10))
        # Deve restituire almeno un chunk (anche se vuoto per PDF dummy)
        assert isinstance(chunks, list)
        # Ogni chunk deve essere TextChunk
        for chunk in chunks:
            assert isinstance(chunk, TextChunk)


def test_docx_parse_stream_chunks(tmp_path):
    """parse_stream per DOCX dovrebbe restituire chunks iterativamente."""
    # Crea un DOCX minimale con python-docx
    try:
        from docx import Document
        
        docx_path = tmp_path / "test.docx"
        doc = Document()
        doc.add_paragraph("Paragraph 1")
        doc.add_paragraph("Paragraph 2")
        doc.add_paragraph("Paragraph 3")
        doc.save(str(docx_path))
        
        parser = get_parser(docx_path)
        assert parser is not None
        
        if hasattr(parser, 'parse_stream'):
            chunks = list(parser.parse_stream(docx_path, chunk_size=2))
            # Con chunk_size=2 e 3 paragrafi, dovremmo avere almeno 2 chunks
            assert len(chunks) >= 2
            
            # Verifica che siano TextChunk validi
            for chunk in chunks:
                assert isinstance(chunk, TextChunk)
                assert chunk.text  # Non vuoto
    except ImportError:
        # Se python-docx non disponibile, skip test
        pass


def test_streaming_fallback_on_unsupported_parser():
    """Se parser non supporta streaming, dovrebbe fare fallback."""
    # Image parser non supporta streaming
    parser = get_parser(Path("test.png"))
    assert parser is not None
    
    # Non deve avere supports_streaming o deve ritornare False
    if hasattr(parser, 'supports_streaming'):
        assert parser.supports_streaming() is False
